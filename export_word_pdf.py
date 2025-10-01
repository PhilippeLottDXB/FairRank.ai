# -*- coding: utf-8 -*-
"""
Created on Wed Sep  3 21:09:15 2025

@author: Admin
"""
import io
import os
import tempfile
import subprocess
from datetime import date
from PyPDF2 import PdfMerger
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_docx_bytes(title, subtitle, highlights=None, table_rows=None):
    doc = Document()

    h = doc.add_heading(title or "Quarterly Performance Snapshot", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(subtitle or f"Generated on {date.today().isoformat()}")
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if highlights:
        doc.add_paragraph("Highlights:")
        for line in highlights.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line, style="List Bullet")

    if table_rows:
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Previous"
        hdr[2].text = "Current"
        for row in table_rows:
            r = table.add_row().cells
            r[0].text, r[1].text, r[2].text = row

    doc.add_paragraph()
    f = doc.add_paragraph("Prepared by: FairRank.ai")
    for run in f.runs:
        run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def try_export_pdf(docx_bytes: bytes) -> tuple[bytes | None, str | None]:
    """
    Attempt PDF export using either:
      A) docx2pdf (Word required), or
      B) LibreOffice headless (soffice).
    Returns (pdf_bytes, method_used_or_error)
    """
    # Write the DOCX to a temp file because both methods operate on files.
    with tempfile.TemporaryDirectory() as td:
        docx_path = os.path.join(td, "report.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        # --- Method A: docx2pdf ---
        try:
            from docx2pdf import convert  # type: ignore
            pdf_path = os.path.join(td, "report_docx2pdf.pdf")
            convert(docx_path, pdf_path)  # raises if Word missing/unusable
            with open(pdf_path, "rb") as f:
                return f.read(), "docx2pdf"
        except Exception as e_a:
            last_error_a = str(e_a)

    # If we get here, no PDF was produced
    return None, f"PDF export unavailable. docx2pdf error: {last_error_a if 'last_error_a' in locals() else 'n/a'}"


def concatenate_pdfs(pdf_files: list[bytes]) -> bytes:
    """
    Concatenate multiple PDFs (given as bytes) into a single PDF.
    Returns the merged PDF as bytes.
    """
    merger = PdfMerger()
    for pdf_bytes in pdf_files:
        merger.append(io.BytesIO(pdf_bytes))
    output_buf = io.BytesIO()
    merger.write(output_buf)
    merger.close()
    output_buf.seek(0)
    return output_buf.getvalue()


if __name__ == "__main__":
    
    st.title("📝 Build & Download a DOCX (and PDF) Report")
    
    with st.form("inputs"):
        col1, col2 = st.columns(2)
        with col1:
            title = st.text_input("Title", "Quarterly Performance Snapshot")
        with col2:
            subtitle = st.text_input("Subtitle", f"Generated on {date.today().isoformat()}")
    
        st.markdown("**Highlights (one per line):**")
        highlights = st.text_area(
            "Highlights",
            "Revenue up 12% QoQ\nChurn down 0.8 percentage points\nNew markets opened in KSA and Qatar",
            height=120,
            label_visibility="collapsed",
        )
    
        st.markdown("**Table (CSV with 3 columns: Metric, Previous, Current)**")
        table_csv = st.text_area(
            "CSV table",
            "Revenue (USD M),8.9,10.0\nActive Users (k),142,158\nNet Retention,108%,112%",
            height=120,
            label_visibility="collapsed",
        )
    
        create_btn = st.form_submit_button("Create report")
    
    if create_btn:
        # Parse table CSV (super lightweight parsing)
        rows = []
        for line in table_csv.splitlines():
            parts = [p.strip() for p in line.split(",")]
            if len(parts) == 3 and any(parts):
                rows.append(parts)
    
        docx_buf = build_docx_bytes(title, subtitle, highlights, rows)
    
        st.success("DOCX report created.")
        st.download_button(
            "⬇️ Download .docx",
            data=docx_buf.getvalue(),
            file_name="report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    
        # Try to create a PDF as a bonus (if environment supports it)
        with st.spinner("Attempting PDF export…"):
            pdf_bytes, method_info = try_export_pdf(docx_buf.getvalue())
        if pdf_bytes:
            st.info(f"PDF generated via **{method_info}**.")
            st.download_button(
                "⬇️ Download .pdf",
                data=pdf_bytes,
                file_name="report.pdf",
                mime="application/pdf",
            )
    
    st.caption("Tip: On Streamlit Community Cloud, LibreOffice/Word aren't available by default. "
               "Deploy on a VM/container where you can install LibreOffice (or run on your own machine).")
    
